import os
import sys
import subprocess

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Instalando python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(37, 99, 235) # Accent color (blue)

def generate_report():
    print("Generando Informe Técnico...")
    doc = docx.Document()
    
    # Titulo
    title = doc.add_heading('Informe Técnico: Plataforma de Alto Impacto Electro10', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Fusión de Data Science, Neuroeducación y TIC').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Autor: Dudbil Olvasada Pabón Riaño\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Arquitectura
    add_heading(doc, '1. Arquitectura de Alto Impacto (Neurociencia + TIC)', 1)
    p = doc.add_paragraph('La plataforma Electro10 ha sido rediseñada para ser mucho más que un repositorio de contenidos. Es un ')
    p.add_run('Ecosistema de Aprendizaje Adaptativo').bold = True
    p.add_run(' que fusiona:')
    doc.add_paragraph('• Neuroeducación: Implementación de consejos dinámicos basados en el tiempo de estudio (curva de atención y fatiga cognitiva) y retención mediante la Ley de Ebbinghaus.', style='List Bullet')
    doc.add_paragraph('• Gamificación Social: Simulación de conexiones de pares para incentivar el aprendizaje social.', style='List Bullet')
    doc.add_paragraph('• Data Science y Analítica Educativa: Un dashboard de administrador ("Control Maestro") que calcula en tiempo real el Índice de Retención Cognitiva, la Carga Cognitiva y una Tasa Predictiva de Abandono (Churn), permitiendo intervención temprana.', style='List Bullet')
    doc.add_paragraph('• Diagnóstico de Presaberes: Un examen evaluativo estricto con un umbral de aprobación del 65%. Analiza el dominio del estudiante en: Cálculo Diferencial, Cálculo Integral, Ecuaciones Diferenciales, Cálculo Vectorial y Álgebra Lineal, generando un Gráfico de Radar Diagnóstico con consejos de nivelación personalizados.', style='List Bullet')

    # 2. Diseño de Alta Gama (Stitch / Material 3)
    add_heading(doc, '2. Arquitectura Visual (Material Design 3 / Stitch Style)', 1)
    doc.add_paragraph('El sistema adopta una estética Premium basada en los principios de diseño de alta gama "Stitch with Google" / Material 3. La interfaz utiliza Glassmorphism, sombreados fluidos y una paleta de colores reactiva al entorno (Light/Dark mode) con acentos dinámicos (Google Blue M3).')

    # 3. GitHub Pages
    add_heading(doc, '3. Guía de Despliegue en GitHub Pages', 1)
    doc.add_paragraph('La plataforma está desarrollada en Vanilla JS, HTML y CSS (Arquitectura Frontend-only sin servidor local requerído), lo que la hace perfecta para GitHub Pages (Hosting gratuito, rápido e infinito).')
    
    doc.add_paragraph('Paso 1: Inicializar Git', style='List Number')
    doc.add_paragraph('Abre la terminal en la carpeta del proyecto y ejecuta:\n git init\n git add .\n git commit -m "Primera versión Electro10"')
    
    doc.add_paragraph('Paso 2: Crear Repositorio en GitHub', style='List Number')
    doc.add_paragraph('Ve a github.com, crea un nuevo repositorio público y vincula tu repositorio local ejecutando los comandos que te proporciona GitHub (git remote add origin ...).')
    
    doc.add_paragraph('Paso 3: Subir código (Push)', style='List Number')
    doc.add_paragraph('git push -u origin main')
    
    doc.add_paragraph('Paso 4: Activar GitHub Pages', style='List Number')
    doc.add_paragraph('En GitHub, ve a Settings > Pages. Selecciona la rama "main" y la carpeta "/root". Haz clic en Save. En 2 minutos tu plataforma estará en línea en https://[tu-usuario].github.io/[tu-repo].')

    # 4. Firebase Backend
    add_heading(doc, '4. Configuración del Backend en Firebase', 1)
    doc.add_paragraph('El "Cerebro" de la plataforma vive en Firebase, brindando Autenticación y Base de Datos en Tiempo Real (Firestore).')
    
    doc.add_paragraph('Paso 1: Crear el Proyecto', style='List Number')
    doc.add_paragraph('Entra a console.firebase.google.com y crea un proyecto llamado "Electro10". Desactiva Google Analytics si no es estrictamente necesario para ahorrar cuota.')
    
    doc.add_paragraph('Paso 2: Activar Authentication', style='List Number')
    doc.add_paragraph('Ve a Build > Authentication. Habilita los proveedores "Correo electrónico/contraseña" y "Google".')
    
    doc.add_paragraph('Paso 3: Activar Firestore Database', style='List Number')
    doc.add_paragraph('Ve a Build > Firestore Database y crea una base de datos en modo producción. Configura las reglas de seguridad básica.')
    
    doc.add_paragraph('Paso 4: Variables de Entorno (config/firebase.config.js)', style='List Number')
    doc.add_paragraph('En la configuración del proyecto de Firebase (Rueda de engranaje), registra una App Web. Copia el objeto firebaseConfig y pégalo en el archivo local "config/firebase.config.js". Asegúrate de NO subir este archivo públicamente a GitHub si quieres seguridad total, aunque Firestore rules protege los datos reales.')

    # 5. Recomendaciones
    add_heading(doc, '5. Recomendaciones de Escalabilidad (Futuro)', 1)
    doc.add_paragraph('1. Machine Learning Nativo: Integrar TensorFlow.js para que el "ElectroBOT IA" aprenda el estilo del usuario y recomiende ejercicios específicos sin llamar a APIs externas.')
    doc.add_paragraph('2. PWA (Progressive Web App): Agregar un archivo manifest.json y un Service Worker para que los estudiantes puedan instalar Electro10 en sus celulares como una aplicación nativa y acceder offline a los talleres.')
    doc.add_paragraph('3. Analítica Predictiva Avanzada: Exportar los datos de Firestore periódicamente a Google BigQuery o Python local para entrenar modelos de Random Forest que determinen exactamente qué temas de matemáticas provocan mayor deserción.')

    doc.add_page_break()
    doc.add_paragraph('Generado automáticamente por Antigravity (Sistema Avanzado de Código).')

    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Informe_Tecnico_Electro10.docx')
    doc.save(output_path)
    print(f"Informe generado exitosamente en: {output_path}")

if __name__ == "__main__":
    generate_report()
